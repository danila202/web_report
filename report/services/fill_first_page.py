from report.models import Application, Fence, Marking, Test


from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt


doc = Document('static/empty_report.docx')
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)


class FistPage:

    @classmethod
    def fill_customer(cls):
        p = doc.paragraphs[16]
        fence = cls.get_fence()
        queryset = Application.objects.select_related('customer', 'fence').filter(fence=fence)
        run = p.add_run(f'{queryset[0].customer.name}, {queryset[0].customer.address} ')
        cls.custom_style(run)

    @classmethod
    def fill_fence(cls):
        p = doc.paragraphs[17]
        fence = cls.get_fence()
        run = p.add_run(f'{fence.marking}')
        cls.custom_style(run)

    @classmethod
    def fill_manufacturer(cls):
        fence = cls.get_fence()
        queryset = Application.objects.select_related('customer', 'fence').filter(fence=fence)
        p = doc.paragraphs[18]
        run = p.add_run(f'{fence.manufacturer}, {queryset[0].customer.address}')
        cls.custom_style(run)

    @classmethod
    def fill_date_of_receipt(cls):
        fence = cls.get_fence()
        date = Application.objects.select_related('fence').filter(fence=fence)
        p = doc.paragraphs[24]
        date_receipt = date[0].receipt_date
        run = p.add_run(f'{date_receipt.day}.{date_receipt.month}.{date_receipt.year} г')
        cls.custom_style(run)

    @classmethod
    def fill_date_of_test(cls):
        fence = cls.get_fence()
        tests = Test.objects.select_related('protocol__fence').filter(protocol__fence=fence)
        date_start = tests[0].data_start
        date_end = tests[0].data_end
        p = doc.paragraphs[25]
        run = p.add_run(f'{date_start.day}.{date_start.month}.{date_start.year}-'
                        f'{date_end.day}.{date_end.month}.{date_end.year} г')
        cls.custom_style(run)

    @classmethod
    def execute(cls):
        cls.fill_customer()
        cls.fill_fence()
        cls.fill_manufacturer()
        cls.fill_date_of_receipt()
        cls.fill_date_of_test()
        doc.save('static/empty_report1.docx')


    @staticmethod
    def get_fence():
        fence = Fence.objects.get(marking=Marking.objects.last().marking)
        return fence

    @classmethod
    def custom_style(cls,run):
        run.font.size = Pt(14)
        run.font.underline = WD_UNDERLINE.DOUBLE
        doc.save('static/empty_report1.docx')





