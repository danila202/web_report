from docx.shared import Pt

from report.models import Element, Fence, TestParameter, Transport

from .fill_first_page import FistPage, doc


class Table:
    all_table = doc.tables
    result = TestParameter.objects.select_related('test__protocol__fence').filter(
        test__protocol__fence__id=FistPage.get_fence().id)

    @classmethod
    def fill_table1_and_table2(cls):
        fence = FistPage.get_fence()
        elements_query = list(Element.objects.filter(elementmaterialfence__fence__marking=fence.marking))
        table1 = cls.all_table[0].rows[2]
        table2 = cls.all_table[1].rows[2]
        table1.cells[0].text, table2.cells[0].text = fence.marking, fence.marking

        for i in range(1, 4):
            table1.cells[i].text = f'{elements_query[i-1].name} {elements_query[i-1].marking}'
            table2.cells[i].text = f'{elements_query[i-1].name} {elements_query[i-1].marking}'

        cls.save_file()

    @classmethod
    def fill_table3(cls):
        fence_retention = Fence.objects.get(marking=FistPage.get_fence().marking).retention_capacity

        table3 = cls.all_table[2]
        row1, row2, row3, row4 = table3.rows[1], table3.rows[2], table3.rows[3], table3.rows[4]
        row1.cells[1].text, row2.cells[1].text, row3.cells[1].text, row4.cells[1].text = 'кДж', 'м', 'м', 'кт'

        #Заполенение 1 строки
        ud = cls.result.get(parameter_id="Удерживающая способ")
        row1.cells[2].text = f'{fence_retention}    ({ud.actual_value} кДж)'
        row1.cells[3].text = f'{fence_retention}    ({ud.actual_value} кДж)'
        row1.cells[4].text = f'{ud.actual_value - ud.actual_value}%'

        #Заполнение 2 строки
        o = cls.result.filter(parameter_id="Остаточный прогиб")
        if o[0].test.protocol.kind == 'виртуальный':
            row2.cells[3].text = f'{o[0].actual_value}'
            row2.cells[2].text = f'{o[1].actual_value}'
        else:
            row2.cells[2].text = f'{o[0].actual_value}'
            row2.cells[3].text = f'{o[1].actual_value}'
        row2.cells[4].text = f'{round(o[0].actual_value - o[1].actual_value,2)}%'

        #Заполнение 3 строчки
        work_width = cls.result.filter(parameter_id="Рабочая ширина")
        if work_width[0].test.protocol.kind == 'виртуальный':
            row3.cells[3].text = f'{work_width[0].actual_value}'
            row3.cells[2].text = f'{work_width[1].actual_value}'
        else:
            row3.cells[2].text = f'{work_width[0].actual_value}'
            row3.cells[3].text = f'{work_width[1].actual_value}'
        row3.cells[4].text = f'{round(work_width[0].actual_value - work_width[1].actual_value,2)}%'

        # Заполнение 4 строчки
        index = cls.result.get(parameter_id='И автобуса')
        row4.cells[2].text = f'{index.actual_value}'
        row4.cells[3].text = f'{index.actual_value}'
        row4.cells[4].text = f'{round(index.actual_value - index.actual_value, 2)}%'

        cls.save_file()

    @classmethod
    def fill_table4(cls):
        table4 = cls.all_table[3]
        row3, row6, row10 = table4.rows[2], table4.rows[5], table4.rows[9]

        # Заполнение 3 строки
        validation = cls.result.get(parameter_id='Валидация')
        if validation.actual_value !=0:
            row3.cells[1].text = f'{validation.actual_value}%'
            row3.cells[2].text = 'Соответствует'

        # Заполнение 6 строки
        node = cls.result.get(parameter_id='УзлыКрепления')
        if node.actual_value != 0:
            row6.cells[1].text = f'{node.actual_value}'
            row6.cells[2].text = 'Соответствует'

        # Заполнение 10 строки

        alteration = cls.result.get(parameter_id="ИзменениеКонстр")
        row10.cells[1].text = f'{alteration.actual_value}'
        row10.cells[2].text = 'Соответствует'

        cls.save_file()



    @classmethod
    def fill_table5(cls):
        ts_data = Transport.objects.filter(test__protocol__fence__id=FistPage.get_fence().id)
        table5 = cls.all_table[4]
        row3, row4, row5, row6, row7, row8, row9 = table5.rows[2],table5.rows[3],table5.rows[4],table5.rows[5],\
        table5.rows[6], table5.rows[7], table5.rows[8]

        row3.cells[1].text, row3.cells[2].text = str(ts_data[0].length), str(ts_data[1].length)
        row4.cells[1].text, row4.cells[2].text = str(ts_data[0].width), str(ts_data[1].width)
        row5.cells[1].text, row5.cells[2].text = str(ts_data[0].track), str(ts_data[1].track)
        row6.cells[1].text, row6.cells[2].text = str(ts_data[0].mass), str(ts_data[1].mass)
        row7.cells[1].text, row7.cells[2].text = str(ts_data[0].load_po), str(ts_data[1].load_po)
        row8.cells[1].text, row8.cells[2].text = str(ts_data[0].load_zo), str(ts_data[1].load_zo)
        row9.cells[1].text, row9.cells[2].text = str(ts_data[0].height), str(ts_data[1].height)

        p_bus = doc.paragraphs[50]
        p_car = doc.paragraphs[52]
        summ_bus = ts_data[0].width + (0.16*ts_data[0].length) + (0.22*20)
        summ_car = ts_data[1].width + (0.16*ts_data[1].length) + (0.22*10)
        run_bus = p_bus.add_run(f'{ts_data[0].width} + 0.16*{ts_data[0].length} + 0.22*20 = {summ_bus}м.')
        run_cur = p_car.add_run(f'{ts_data[1].width} + 0.16*{ts_data[1].length} + 0.22*10 = {summ_car}м.')
        run_bus.font.size = Pt(14)
        run_cur.font.size = Pt(14)


        cls.save_file()


    @classmethod
    def fill_table6(cls):
        table6 = cls.all_table[5]
        # Заполняем 4
        row4 = table6.rows[3]
        trajectory = cls.result.filter(parameter_id='Траектория')
        if trajectory[0].actual_value == 0 and trajectory[1].actual_value == 0 :
            row4.cells[1].text = 'Не изменило своей траектории движения'
            row4.cells[2].text = 'Не изменило своей траектории движения'

        # Заполняем 5
        speed = cls.result.filter(parameter_id='Скорость')
        row5 = table6.rows[4]
        row5.cells[1].text, row5.cells[2].text = f'{speed[0].actual_value}%', f'{speed[1].actual_value}%'

        # Заполняем 6
        row6 = table6.rows[5]
        tipping = cls.result.filter(parameter_id='Опрокидывание')
        if tipping[0].actual_value == 0 and tipping[1].actual_value == 0:
            row6.cells[1].text = 'Опрокидывания не произошло, автобус не пересек линию ограждения'
            row6.cells[2].text = 'Опрокидывания не произошло, легковой автомобиль не переехал ограждение'

        # Заполняем 7
        row7 = table6.rows[6]
        damage = cls.result.filter(parameter_id='Повреждения')
        if damage[0].actual_value == 0 and damage[1].actual_value == 0:
            row7.cells[1].text = 'Автобус не получил серьезных повреждений'
            row7.cells[2].text = 'Легковой автомобиль не получил серьезных повреждений'

        row8 = table6.rows[7]
        ts_data = Transport.objects.filter(test__protocol__fence__id=FistPage.get_fence().id)
        summ_bus = ts_data[0].width + (0.16 * ts_data[0].length) + (0.22 * 20)
        summ_car = ts_data[1].width + (0.16 * ts_data[1].length) + (0.22 * 10)
        row8.cells[1].text = f'K = {summ_bus}м. B = 20м.'
        row8.cells[2].text = f'K = {summ_car}м. B = 10м.'
        p = doc.paragraphs[35]
        run = p.add_run(f'марки {FistPage.get_fence().marking} методом наезда автобусом массой {ts_data[0].mass} т '
                        'со скоростью 67 км/ч и '\
         f'методом наезда легковым автомобилем массой {ts_data[1].mass} т со скоростью 90 км/ч.')

        run.font.size = Pt(14)
        cls.save_file()


    @classmethod
    def fill_table7(cls):
        table7 = cls.all_table[6]

        #Заполняем 3
        row3 = table7.rows[2]
        bus = cls.result.get(parameter_id='И автобуса')
        if bus.actual_value < 1.1:
            row3.cells[1].text = str(bus.actual_value)
        else:
            row3.cells[1].text = ''

        car = cls.result.get(parameter_id='И автобуса')
        if car.actual_value < 1:
            row3.cells[2].text = str(car.actual_value)
        else:
            row3.cells[2].text = ''

        row4 = table7.rows[3]
        avg_kf = cls.result.filter(parameter_id='СрКоэф СВР')
        min_kf = cls.result.filter(parameter_id='МинКоэфСВР')

        if avg_kf[0].actual_value >= 0.9 and min_kf[0].actual_value>=0.8:
            row4.cells[1].text = f'{avg_kf[0].actual_value}/{min_kf[0].actual_value}'

        if avg_kf[1].actual_value >= 0.9 and min_kf[1].actual_value>=0.8:
            row4.cells[2].text = f'{avg_kf[1].actual_value}/{min_kf[1].actual_value}'

        cls.save_file()

    @classmethod
    def fill_table8(cls):
        tabl8 = cls.all_table[7]
        rupture = cls.result.filter(parameter_id='Разрыв')
        row4 = tabl8.rows[3]
        if rupture[0].actual_value == 0:
            row4.cells[1].text = 'Разрыва направляющей балки не произошло'
            row4.cells[2].text = 'Разрыва направляющей балки не произошло'

        cls.save_file()

    @classmethod
    def fill_table9(cls):
        table9 = cls.all_table[8]

        row3, row4, row5 = table9.rows[2], table9.rows[3], table9.rows[4]
        row3.cells[1].text, row4.cells[1].text, row5.cells[1].text = 'кДж', 'м', 'м'

        fence_retention = FistPage.get_fence().retention_capacity
        ud = cls.result.get(parameter_id="Удерживающая способ")
        row3.cells[2].text = f'{fence_retention}    ({ud.actual_value} кДж)'
        row3.cells[3].text = f'{fence_retention}    ({ud.actual_value} кДж)'

        ls = cls.result.filter(parameter_id='Остаточный прогиб')

        if ls[0].test_id % 2 == 0:
            row4.cells[2].text = str(ls[1].actual_value)
            row4.cells[3].text = str(ls[0].actual_value)
        else:
            row4.cells[2].text = str(ls[0].actual_value)
            row4.cells[3].text = str(ls[1].actual_value)

        work_width = cls.result.filter(parameter_id="Рабочая ширина")


        if work_width[0].test_id % 2 == 0:
            row5.cells[2].text = str(work_width[1].actual_value)
            row5.cells[3].text = str(work_width[0].actual_value)
        else:
            row5.cells[2].text = str(work_width[0].actual_value)
            row5.cells[3].text = str(work_width[1].actual_value)


        cls.save_file()

    @classmethod
    def fill_table10(cls):
        table10 = cls.all_table[9]

        row3, row4, row5, row6 = table10.rows[2], table10.rows[3], table10.rows[4], table10.rows[5]
        row3.cells[1].text, row4.cells[1].text, row5.cells[1].text, row6.cells[1].text = 'шт.', 'шт.', 'шт.', 'м'

        rack = cls.result.filter(parameter_id='Кол-во деформ стоек')
        row3.cells[2].text, row3.cells[3].text = str(rack[0].actual_value), str(rack[1].actual_value)

        beam = cls.result.filter(parameter_id='Кол-во деформ балок')
        row4.cells[2].text, row4.cells[3].text = str(beam[0].actual_value), str(beam[1].actual_value)

        console = cls.result.filter(parameter_id='Кол-во деформ конс')
        row5.cells[2].text, row5.cells[3].text = str(console[0].actual_value), str(console[1].actual_value)

        length = cls.result.filter(parameter_id='Длина взаимодействия')
        row6.cells[2].text, row6.cells[3].text = str(length[0].actual_value), str(length[1].actual_value)

        cls.save_file()



    @classmethod
    def execute(cls):
        cls.fill_table1_and_table2()
        cls.fill_table3()
        cls.fill_table4()
        cls.fill_table5()
        cls.fill_table6()
        cls.fill_table7()
        cls.fill_table8()
        cls.fill_table9()
        cls.fill_table10()

    @staticmethod
    def save_file():
        doc.save('static/empty_report1.docx')



