

from django.shortcuts import render,HttpResponse
from .forms import FencingForm
from .models import  Fence, Element, Customer

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt, Mm
import pypandoc


doc = Document('static/empty_report.docx')


def home(request):
    if request.method == 'POST':
        form = FencingForm(request.POST)
        if form.is_valid():
            if Fence.objects.filter(marking=request.POST.get('marking')).exists():
                fill_customer(request)

                context = pypandoc.convert_file('static/empty_report.docx', 'html', format='docx')
                return render(request, 'html/docx.html', {'context': context})

            else:
                return HttpResponse('Такого ограждения нет')

    else:
        form = FencingForm()

    return render(request, 'html/home.html', {'form': form})


def fill_customer(request):
    p = doc.paragraphs[16]
    queryset = Customer.objects.get(id=1)
    run = p.add_run(f'{queryset.name}, {queryset.address} ')
    run.font.size = Pt(14)
    run.font.underline = WD_UNDERLINE.DOUBLE
    doc.save('static/empty_report.docx')


def fill_table1(request):
    ...








